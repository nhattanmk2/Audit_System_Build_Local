import os
import re
import copy
import zipfile
import io
from docx import Document
from typing import List, Dict, Any

def split_reports(template_path: str, output_dir: str):
    """
    Splits the main report into individual reports for each system.
    Using XML-level deep copy to preserve formatting / merges / styles exactly as in the original.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document(template_path)
    
    # Identify systems and their content based on headers like "1.1. Name"
    # We ignore anything before the first system header (like TOC)
    system_pattern = re.compile(r'^([123]\.\d+)\.\s+(.*)$')
    
    # Map elements directly from doc.element.body to their correct Python wrappers
    # doc.paragraphs only contains <w:p>, doc.tables only <w:tbl>
    # We need the full list in order to group them.
    body_elements = list(doc.element.body)
    
    systems = []
    current_system = None
    
    for i, element in enumerate(body_elements):
        # We only care about paragraphs and tables
        if element.tag.endswith('p'):
            text = "".join(element.itertext()).strip()
            match = system_pattern.match(text)
            if match:
                if current_system:
                    systems.append(current_system)
                
                current_system = {
                    "id": match.group(1),
                    "name": match.group(2).strip(),
                    "full_title": text,
                    "elements": [element] # Start with the title paragraph itself
                }
                continue
        
        elif element.tag.endswith('tbl'):
            pass # Tables are added to current_system below
        
        else:
            # Skip other elements like sectPr if they aren't part of a system's block
            continue

        if current_system:
            current_system["elements"].append(element)

    if current_system:
        systems.append(current_system)

    # Now create individual documents
    generated_files = []
    for sys in systems:
        # Create a new document using the SAME template to keep styles/headers/margins
        new_doc = Document(template_path)
        
        # CLEAR the body elements (<w:p> and <w:tbl>)
        # but keep other properties like section properties if they exist
        body = new_doc.element.body
        for element in list(body):
            if element.tag.endswith(('p', 'tbl')):
                body.remove(element)
        
        # APPEND copied elements from the original document
        for elem in sys["elements"]:
            # Perform a deep copy of the XML element
            new_elem = copy.deepcopy(elem)
            
            # If it's a table, we need to clear the result column (cột 1)
            # but ONLY for rows > 0 (skip header)
            if new_elem.tag.endswith('tbl'):
                # Using lower-level XML manipulation to clear cell text
                rows = new_elem.findall(qn('w:tr'))
                if rows:
                    for r_idx, row in enumerate(rows):
                        if r_idx == 0: continue # Skip header
                        
                        cells = row.findall(qn('w:tc'))
                        # Clear column 1 (index 1) for both 2-col and 3-col tables
                        if len(cells) >= 2:
                            result_cell = cells[1]
                            # Clear all content in the cell
                            for p in result_cell.findall(qn('w:p')):
                                # We keep the paragraph element but clear its text (r elements)
                                for r in p.findall(qn('w:r')):
                                    p.remove(r)

            body.append(new_elem)

        # Save the file
        # Sanitize filename
        safe_name = re.sub(r'[\\/*?:"<>|]', "_", sys["name"])
        file_path = os.path.join(output_dir, f"{safe_name}.docx")
        new_doc.save(file_path)
        generated_files.append(file_path)

    return generated_files

def fill_audit_results(template_file: str, audit_results: List[Dict[str, Any]], output_path: str, llm=None):
    """
    Fills the audit results (PASS/FAIL) into a system-specific docx template.
    If llm is provided, use it to intelligently map criteria names to audit results.
    """
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    import json

    doc = Document(template_file)
    
    # 1. Collect all criteria names from the Word tables
    all_rows = []
    for table_idx, table in enumerate(doc.tables):
        if len(table.columns) < 2:
            continue
        
        header_text = table.rows[0].cells[0].text.lower()
        if "hạng mục" not in header_text and "nội dung" not in header_text:
            continue

        for row_idx, row in enumerate(table.rows[1:], 1):
            criteria_text = row.cells[0].text.strip()
            if criteria_text:
                all_rows.append({
                    "table_idx": table_idx,
                    "row_idx": row_idx,
                    "criteria": criteria_text
                })

    if not all_rows:
        doc.save(output_path)
        return output_path

    # 2. Use AI to map the Word criteria to audit results
    if llm and audit_results:
        # Prepare data for LLM
        # We only need enough info for matching: rule_id and title
        results_summary = []
        for r in audit_results:
            # We look in r["result"] or r depending on structure
            res_data = r.get("result", r)
            results_summary.append({
                "rule_id": res_data.get("rule_id", "N/A"),
                "title": res_data.get("title", "N/A")
            })

        # Batch criteria mapping if there are many rows to avoid context limits
        # However, typically a single system report has 20-50 rows, fitting in one prompt.
        prompt = ChatPromptTemplate.from_messages([
            ("system", """
Bạn là một trợ lý kỹ thuật chuyên về đối soát an ninh thông tin.

NHIỆM VỤ:
So khớp danh sách các 'Tiêu chí trong Word' (Tiếng Việt) với danh sách 'Kết quả Audit' (Thường là Rule ID và Tiêu đề Tiếng Anh/Kỹ thuật).

YÊU CẦU:
1. Xác định Tiêu chí Word tương ứng với Rule ID nào.
2. Trả về kết quả dưới dạng JSON object: {{"criteria_index": "rule_id", ...}}
3. Chỉ trả về JSON, không giải thích gì thêm.
4. Nếu không khớp được bản ghi nào, bỏ qua hoặc gán "NONE".
            """),
            ("user", """
DANH SÁCH TIÊU CHÍ TRONG WORD:
{word_criteria}

DANH SÁCH KẾT QUẢ AUDIT:
{audit_results}

HÃY THỰC HIỆN SO KHỚP VÀ TRẢ VỀ JSON:
            """)
        ])

        # Convert to string list for prompt
        word_criteria_str = "\n".join([f"{i}. {item['criteria']}" for i, item in enumerate(all_rows)])
        audit_results_str = json.dumps(results_summary, ensure_ascii=False, indent=2)

        chain = prompt | llm | StrOutputParser()
        
        try:
            mapping_raw = chain.invoke({
                "word_criteria": word_criteria_str,
                "audit_results": audit_results_str
            })
            
            # Extract JSON
            mapping_raw = mapping_raw.replace("```json", "").replace("```", "").strip()
            mapping = json.loads(mapping_raw)
            
            # Map rule_id to status for easy lookup
            status_map = {}
            for r in audit_results:
                res_data = r.get("result", r)
                status_map[res_data.get("rule_id")] = res_data.get("compliance_status", "SKIP")

            # 3. Apply the results back to the document
            for i, row_info in enumerate(all_rows):
                rule_id = mapping.get(str(i))
                if rule_id and rule_id != "NONE" and rule_id in status_map:
                    status = status_map[rule_id].upper()
                    # Apply to table
                    target_table = doc.tables[row_info["table_idx"]]
                    target_row = target_table.rows[row_info["row_idx"]]
                    target_row.cells[1].text = status
        
        except Exception as e:
            print(f"Error during AI matching in Word fill: {e}")
            # Fallback to basic matching logic if AI fails
            pass 

    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    # Test splitting
    template = r"e:\CIS_Audit_Project\CIS_Data\9_Báo_cáo_kết_quả_kiểm_định_an_toàn_hệ_thống_Fix_done.docx"
    out = r"e:\CIS_Audit_Project\CIS_Data\Reports_Template"
    # split_reports(template, out)
