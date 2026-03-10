from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

def generate_cmc_report(audit_data, target_server="N/A"):
    """
    audit_data: list of dicts {
        "param_name": str,
        "actual_value": str,
        "result": "PASS" | "FAIL",
        "remediation": str
    }
    """
    doc = Document()

    # --- Header / Title ---
    title = doc.add_heading('BÁO CÁO KIỂM ĐỊNH AN TOÀN THÔNG TIN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Section I: Tổng quan ---
    doc.add_heading('I. TỔNG QUAN HỆ THỐNG', level=1)
    
    table_info = doc.add_table(rows=2, cols=2)
    table_info.style = 'Table Grid'
    table_info.cell(0, 0).text = 'Tên máy chủ:'
    table_info.cell(0, 1).text = target_server
    table_info.cell(1, 0).text = 'Thời gian quét:'
    table_info.cell(1, 1).text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    doc.add_paragraph()
    
    total = len(audit_data)
    passed = len([r for r in audit_data if r["result"] == "PASS"])
    failed = total - passed
    
    p = doc.add_paragraph()
    p.add_run(f"Kết quả kiểm tra: ").bold = True
    p.add_run(f"Tổng số {total} tham số. ")
    p.add_run(f"Đạt: {passed}").font.color.rgb = None # Green color would be nice but simple text for now
    p.add_run(f", Không đạt: {failed}").font.color.rgb = None

    # --- Section II: Chi tiết kết quả ---
    doc.add_heading('II. CHI TIẾT KẾT QUẢ KIỂM TRA', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Danh mục / Tham số'
    hdr_cells[2].text = 'Kết quả'
    hdr_cells[3].text = 'Khuyến nghị / Remediation'

    for idx, item in enumerate(audit_data, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item["param_name"]
        row_cells[2].text = item["result"]
        row_cells[3].text = item.get("remediation", "")

    # Save to buffer
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream
